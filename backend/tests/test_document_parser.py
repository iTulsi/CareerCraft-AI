from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.services.document_parser import (
    DocumentParseError,
    MAX_FILE_SIZE_BYTES,
    extract_text,
)
from app.services.section_parser import detect_sections


client = TestClient(app)


def _build_docx() -> bytes:
    document = Document()
    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(
        "Python developer building reliable AI applications."
    )
    document.add_heading("Technical Skills", level=1)
    document.add_paragraph("Python, FastAPI, React, Docker")
    document.add_heading("Projects", level=1)
    document.add_paragraph("CareerCraft AI resume analysis platform.")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extract_text_from_docx() -> None:
    text = extract_text("resume.docx", _build_docx())

    assert "Python developer" in text
    assert "CareerCraft AI" in text


def test_detect_sections() -> None:
    sections = detect_sections(
        """
        Professional Summary
        Python developer building AI systems.
        Technical Skills
        Python, FastAPI, Docker
        Projects
        CareerCraft AI
        """
    )

    assert "Python developer" in sections["summary"]
    assert "FastAPI" in sections["skills"]
    assert "CareerCraft AI" in sections["projects"]


def test_parse_resume_endpoint() -> None:
    response = client.post(
        "/api/resume/parse",
        files={
            "file": (
                "resume.docx",
                _build_docx(),
                (
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"] == "resume.docx"
    assert payload["word_count"] > 5
    assert "skills" in payload["sections"]


def test_rejects_unsupported_file_type() -> None:
    response = client.post(
        "/api/resume/parse",
        files={"file": ("resume.png", b"not-an-image", "image/png")},
    )

    assert response.status_code == 422
    assert "Unsupported file type" in response.json()["detail"]


def test_rejects_empty_text_file() -> None:
    try:
        extract_text("resume.txt", b"")
    except DocumentParseError as exc:
        assert "empty" in str(exc).lower()
    else:
        raise AssertionError("Expected DocumentParseError")


def test_parse_resume_endpoint_rejects_oversized_upload() -> None:
    response = client.post(
        "/api/resume/parse",
        files={
            "file": (
                "resume.txt",
                b"a" * (MAX_FILE_SIZE_BYTES + 1),
                "text/plain",
            )
        },
    )

    assert response.status_code == 422
    assert "exceeds the 5 MB limit" in response.json()["detail"]


def test_normalizes_common_pdf_ligatures() -> None:
    from app.services.document_parser import normalize_extracted_text

    text = normalize_extracted_text(
        "Machine learning \ufb01ne-tuning and \ufb02exible workflows."
    )

    assert text == "Machine learning fine-tuning and flexible workflows."


def test_detects_sections_in_flattened_pdf_text() -> None:
    flattened_text = (
        "TULSI TOMAR PROFESSIONAL SUMMARY Python developer building AI systems. "
        "TECHNICAL SKILLS Python FastAPI Docker PROJECTS CareerCraft AI "
        "EDUCATION B.Tech Computer Science"
    )

    sections = detect_sections(flattened_text)

    assert "Python developer" in sections["summary"]
    assert "FastAPI" in sections["skills"]
    assert "CareerCraft AI" in sections["projects"]
    assert "B.Tech" in sections["education"]
